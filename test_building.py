#!/usr/bin/env python3
"""生成建筑行业日报Word文档 - 测试运行"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

now = datetime.now()
day = f"{now.month}月{now.day}日"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("建筑行业日报")
r.bold = True; r.font.size = Pt(22)
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(f"{now.strftime('%Y年%m月%d日')}")
sr.font.size = Pt(10)
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph("━" * 50)

# 模拟建筑行业新闻
news = [
    ("水利部推进重大水利工程建设，2026年投资规模超万亿",
     "水利部发布2026年重大水利工程项目清单，全年计划新开工45项重大水利工程，总投资规模超过1.2万亿元。其中涉及多个大型水库、引调水及灌区项目，对水利监理工程师需求旺盛。",
     "http://www.mwr.gov.cn/"),
    ("住建部：推进智能建造与建筑工业化协同发展",
     "住建部发布新政策，要求2026年起新开工项目全面推广BIM技术应用，鼓励建筑机器人、装配式建筑等新型建造方式。持有注册监理工程师、一级建造师等证书的技术人才需求持续增长。",
     "https://www.mohurd.gov.cn/"),
    ("国资委部署中央企业高质量发展的重点任务",
     "国资委召开中央企业负责人会议，强调央企要聚焦主责主业，在基础设施建设、水利工程、城市更新等领域发挥骨干作用。会议特别指出要加强项目管理人才队伍建设。",
     "http://www.sasac.gov.cn/"),
    ("多地发布2026年重点建设项目清单",
     "广东、浙江、四川等省份陆续发布2026年重点建设项目清单，总投资合计超过10万亿元。涵盖交通、水利、能源、城市更新等领域，对工程监理、项目管理人才需求旺盛。",
     "https://www.ndrc.gov.cn/"),
    ("三峡集团启动重大水利枢纽工程招标",
     "长江三峡集团公告多个重大水利枢纽工程项目监理招标，涉及大坝建设、水电站改造、生态修复等，总投资规模超百亿元。要求投标单位配备注册监理工程师（水利工程）。",
     "https://www.ctg.com.cn/"),
    ("中国建筑中标多个大型基建项目",
     "中国建筑集团公告中标多个大型基础设施项目，合同总额超过500亿元，涵盖城市轨道交通、高速公路、大型公共建筑等，急需项目总监、总监理工程师等高端管理人才。",
     "http://www.cscec.com/"),
    ("十四五重大工程进入建设高峰期",
     "随着十四五规划中期评估完成，一批重大工程进入全面建设高峰期。国家发改委表示将加快推动重大项目建设，确保形成更多实物工作量。建筑和工程监理行业景气度持续提升。",
     "https://www.ndrc.gov.cn/"),
    ("工程监理行业改革：推行全过程工程咨询",
     "住建部推动工程监理行业转型升级，鼓励监理企业向全过程工程咨询服务延伸。总监理工程师、注册监理工程师等专业人才在新模式下价值进一步提升。",
     "https://www.mohurd.gov.cn/"),
    ("南水北调后续工程加快推进",
     "南水北调中线后续工程及西线工程前期工作加快推进，多个勘察设计及监理标段已开始招标。水利工程监理及项目管理专业人才需求持续扩大。",
     "http://www.nsbd.gov.cn/"),
    ("国企改革深化提升行动推进，建筑央企整合加速",
     "国资委推进新一轮国企改革，多家建筑类央企进行专业化整合，提升市场竞争力。整合后对项目总监、技术负责人等岗位人才需求增加。",
     "http://www.sasac.gov.cn/"),
    ("城市更新行动全面展开",
     "全国多地启动城市更新行动，涉及老旧小区改造、地下管网更新、智慧城市建设等，总投资规模巨大。工程监理和质量安全管理人员需求增加。",
     "https://www.mohurd.gov.cn/"),
    ("绿色建筑标准全面升级",
     "新版《绿色建筑评价标准》正式实施，对建筑节能、低碳、环保提出更高要求。具有中级以上职称、注册监理工程师等资质的专业人才更受青睐。",
     "https://www.mohurd.gov.cn/"),
]

for i, (title, summary, link) in enumerate(news, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {title}")
    r.bold = True; r.font.size = Pt(12)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    p2 = doc.add_paragraph(summary)
    p2.paragraph_format.first_line_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(1)
    
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run(f"来源：{link}")
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

doc.add_paragraph("━" * 50)
footer = doc.add_paragraph(f"整理日期：{now.strftime('%Y-%m-%d %H:%M')}")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

outdir = r"D:\===每日资讯===\建筑行业日报\2026年5月"
os.makedirs(outdir, exist_ok=True)
fname = f"建筑行业日报_{day}.docx"
path = os.path.join(outdir, fname)
doc.save(path)
print(f"已保存：{path} ({os.path.getsize(path)} bytes, {len(news)}条)")
