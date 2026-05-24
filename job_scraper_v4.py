# -*- coding: utf-8 -*-
"""
国企招聘采集 v4
由近及远采集，附件下载，正文Word存档
"""

import os, re, json, ssl, zipfile, shutil, tempfile, hashlib
import subprocess, urllib.request, urllib.parse
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================================
#  Config
# ============================================================
BASE_DIR = r"D:===国聘==="
TMP_DIR = os.path.join(tempfile.gettempdir(), "guopin_v4")
HISTORY_FILE = os.path.join(BASE_DIR, "_history_v4.json")

ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE

USER_PROFILE = {
    "age": 48, "edu": "大专", "major": "工民建",
    "title": "中级职称",
    "certs": ["注册监理工程师(建筑)", "注册监理工程师(水利)", "一级建造师(建筑)"],
    "city": "南昌", "province": "江西",
}

RECRUIT_KW = ["招聘", "招录", "选聘", "诚聘", "公开招", "社会招",
              "人才引进", "招聘公告", "招聘简章", "招聘启事", "招聘通知",
              "岗位招聘", "人员招聘", "人才招聘", "招聘计划",
              "招人", "招募", "招聘方案"]

MATCH_KW = ["监理", "工程", "建筑", "土木", "施工", "土建", "结构",
            "建造师", "工程师", "建设", "基建", "项目", "技术",
            "总监", "总工", "设计", "造价", "预算", "招标",
            "水利", "市政", "路桥", "交通", "质量", "安全",
            "中建", "中铁", "中交", "中水", "城投", "水投", "建工",
            "国企", "岗位", "技术岗", "管理岗", "工程岗"]

EXCLUDE_KW = ["销售", "客服", "前台", "保洁", "保安", "司机",
              "厨师", "服务员", "普工", "操作工", "包装", "搬运",
              "快递", "外卖", "配送", "促销", "导购", "收银"]

# ============================================================
#  Utils
# ============================================================
def log(msg):
    print("[%s] %s" % (datetime.now().strftime("%H:%M:%S"), msg))

def curl_get(url, timeout=20):
    tmp = tempfile.mktemp(suffix=".html")
    cmd = ["curl.exe", "-s", "-L", "--compressed", "-o", tmp, "-m", str(timeout),
           "-H", "User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
           "-H", "Accept: text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
           "-H", "Accept-Language: zh-CN,zh;q=0.9",
           "--insecure"]
    cmd.append(url)
    try:
        subprocess.run(cmd, capture_output=True, timeout=timeout+5)
        if os.path.exists(tmp) and os.path.getsize(tmp) > 100:
            with open(tmp, "r", encoding="utf-8", errors="replace") as f:
                data = f.read()
            os.remove(tmp)
            return data
        try:
            os.remove(tmp)
        except:
            pass
    except:
        try:
            os.remove(tmp)
        except:
            pass
    return ""

def download_file(url, save_path):
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        resp = urllib.request.urlopen(req, context=ctx, timeout=30)
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        with open(save_path, "wb") as f:
            f.write(resp.read())
        return os.path.getsize(save_path)
    except:
        return 0

def is_recruitment(title, content=""):
    text = (title + " " + content[:200]).lower()
    for kw in EXCLUDE_KW:
        if kw in text:
            return False
    for kw in RECRUIT_KW:
        if kw in text:
            return True
    return False

def is_matching(title, content=""):
    text = (title + " " + content[:500]).lower()
    for kw in MATCH_KW:
        if kw in text:
            return True
    return False

def is_eligible(title, content=""):
    if not is_recruitment(title, content):
        return False
    if not is_matching(title, content):
        return False
    return True

def extract_attachments(html, base_url):
    exts = r"\.(pdf|docx?|xlsx?|pptx?|rar|zip|7z)"
    links = set()
    for m in re.finditer(r'href=["\']([^"\']+)["\']', html, re.I):
        href = m.group(1).strip()
        if re.search(exts, href, re.I):
            if href.startswith("http"):
                links.add(href)
            elif href.startswith("/"):
                from urllib.parse import urlparse
                p = urlparse(base_url)
                links.add(f"{p.scheme}://{p.netloc}{href}")
            else:
                from urllib.parse import urljoin
                links.add(urljoin(base_url, href))
    return list(links)

def clean_html_content(html):
    """返回纯文本，去掉脚本样式和广告词"""
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL|re.I)
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL|re.I)
    text = re.sub(r'<[^>]+>', '', html)
    text = re.sub(r'\s+', ' ', text).strip()
    lines = text.split(chr(10))
    clean = []
    ad_words = ["广告", "推广", "扫码", "关注", "订阅", "分享", "点赞",
                "在看", "转发", "收藏", "星标", "设为星标", "点击", "链接"]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if any(k in line for k in ad_words):
            continue
        if len(line) < 4:
            continue
        if re.match(r'^[\d\s\-=*.,，。、！？:;【】()（）]+$', line):
            continue
        clean.append(line)
    return "\n".join(clean)

def strip_urls(text):
    return re.sub(r'https?://[^\s,，。、；;]+', '', text)

# ============================================================
#  History / Dedup
# ============================================================
def load_history():
    try:
        return json.load(open(HISTORY_FILE, "r", encoding="utf-8"))
    except:
        return {"seen": []}

def save_history(h):
    os.makedirs(os.path.dirname(HISTORY_FILE), exist_ok=True)
    cutoff = (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d")
    h["seen"] = [s for s in h["seen"] if s.get("date", "") >= cutoff]
    json.dump(h, open(HISTORY_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=2)

def dedup_key(title, source=""):
    raw = re.sub(r"[^\u4e00-\u9fa5a-zA-Z0-9]", "", (title+source).lower())[:30]
    return hashlib.md5(raw.encode()).hexdigest()

def is_seen(title, source, history):
    key = dedup_key(title, source)
    for s in history["seen"]:
        if s["key"] == key and (datetime.now() - datetime.strptime(s["date"], "%Y-%m-%d")).days < 14:
            return True
    return False

def mark_seen(title, source, url, history):
    history["seen"].append({
        "key": dedup_key(title, source),
        "title": title[:50],
        "source": source,
        "url": url,
        "date": datetime.now().strftime("%Y-%m-%d")
    })

# ============================================================
#  Sources: WeChat via Sogou
# ============================================================
def search_wechat_articles(keyword, page=1):
    kw = urllib.parse.quote(keyword)
    url = "https://weixin.sogou.com/weixin?type=2&query=%s&page=%d&ie=utf8" % (kw, page)
    html = curl_get(url)
    if not html:
        return []
    articles = []
    # Try mp.weixin.qq.com links first
    for m in re.finditer(
        r'<a[^>]*href="(https?://mp\.weixin\.qq\.com[^"\s]+)"[^>]*>(.*?)</a>',
        html, re.DOTALL
    ):
        link = m.group(1)
        title = re.sub(r'<[^>]+>', '', m.group(2)).strip()
        if not title or len(title) < 4:
            continue
        source = ""
        src_m = re.search(r'class="account"[^>]*>([^<]+)', html)
        if src_m:
            source = src_m.group(1).strip()
        articles.append({"title": title, "link": link, "source": source or "微信公众号"})
    if not articles:
        for m in re.finditer(
            r'<h3[^>]*>.*?<a[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
            html, re.DOTALL
        ):
            link = m.group(1)
            title = re.sub(r'<[^>]+>', '', m.group(2)).strip()
            if not title or len(title) < 4:
                continue
            if link.startswith("/"):
                link = "https://weixin.sogou.com" + link
            source = ""
            src_m = re.search(r'class="account"[^>]*>([^<]+)', html[
                max(0, html.find(m.group(0))-200):html.find(m.group(0))+200
            ])
            if src_m:
                source = src_m.group(1).strip()
            articles.append({"title": title, "link": link, "source": source or "微信公众号"})
    return articles

# ============================================================
#  Sources: Government websites
# ============================================================
def scrape_government_site(name, url):
    html = curl_get(url)
    if not html or len(html) < 300:
        return []
    jobs = []
    for m in re.finditer(r'<a[^>]*href=["\']([^"\']+)["\']>(.*?)</a>', html, re.DOTALL):
        link = m.group(1).strip()
        title = re.sub(r'<[^>]+>', '', m.group(2)).strip()
        if not title or len(title) < 5:
            continue
        if link.startswith("//"):
            link = "https:" + link
        elif link.startswith("/"):
            from urllib.parse import urlparse
            p = urlparse(url)
            link = "%s://%s%s" % (p.scheme, p.netloc, link)
        elif not link.startswith("http"):
            from urllib.parse import urljoin
            link = urljoin(url, link)
        jobs.append({"title": title, "link": link, "source": name})
    return jobs

# ============================================================
#  Output generation
# ============================================================
def make_output_dir():
    now = datetime.now()
    return os.path.join(BASE_DIR, "%d年%d月" % (now.year, now.month),
                        "%d月%d日" % (now.month, now.day))

def create_clean_word_doc(title, source, content, att_urls, save_path):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)
    st = doc.styles["Normal"]
    st.font.name = "SimSun"
    st.font.size = Pt(11)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "SimHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("来源：%s | 采集时间：%s" % (source,
        datetime.now().strftime("%Y-%m-%d %H:%M"))).font.size = Pt(9)

    doc.add_paragraph("-" * 50)
    if content:
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Cm(0.5)
                p.paragraph_format.line_spacing = Pt(24)
    if att_urls:
        doc.add_paragraph("")
        doc.add_paragraph("-" * 50)
        doc.add_heading("附件清单", level=3)
        for a in att_urls:
            doc.add_paragraph("* " + os.path.basename(a))
    doc.add_paragraph("")
    doc.add_paragraph("-" * 50)
    footer = doc.add_paragraph("本文件由 AI 自动采集生成")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    return os.path.getsize(save_path)

def process_job(job, history, tmp_dir):
    title = job["title"]
    link = job["link"]
    source = job["source"]
    log("  Processing: %s..." % title[:40])
    article_dir = os.path.join(tmp_dir, dedup_key(title, source))
    os.makedirs(article_dir, exist_ok=True)
    attachments = []
    content = ""
    page_html = curl_get(link)
    if page_html and len(page_html) > 500:
        clean = clean_html_content(page_html)
        content = strip_urls(clean)
        att_urls = extract_attachments(page_html, link)
        for att_url in att_urls:
            att_name = os.path.basename(att_url.split("?")[0])
            if not att_name or len(att_name) < 5:
                ext_m = re.search(r'\.(pdf|docx?|xlsx?|pptx?|rar|zip|7z)', att_url, re.I)
                att_name = "attachment_%s.%s" % (hashlib.md5(att_url.encode()).hexdigest()[:8],
                                                  ext_m.group(1) if ext_m else "dat")
            save_path = os.path.join(article_dir, att_name)
            sz = download_file(att_url, save_path)
            if sz > 100:
                attachments.append(save_path)
                log("    Downloaded: %s (%dKB)" % (att_name, sz//1024))
    else:
        att_urls = []
    docx_path = os.path.join(article_dir, "article.docx")
    create_clean_word_doc(title, source, content,
                          att_urls if 'att_urls' in dir() else [], docx_path)
    mark_seen(title, source, link, history)
    return {"title": title, "source": source, "link": link,
            "docx": docx_path, "attachments": attachments}

# ============================================================
#  Source lists
# ============================================================
WECHAT_ACCOUNTS = [
    "聚才江西招聘求职", "南昌国企招聘", "南昌招聘站",
    "赣鄱人才站", "南昌人才", "南昌招聘象",
    "江西招聘象", "赣鄱人力", "南昌招聘云",
    "南昌招聘", "南昌人才招聘网", "南昌就创人力",
]

LOCAL_GOV_SITES = [
    ("南昌人才招聘网", "https://www.ncrczpw.com/"),
    ("江西省人社厅", "http://rst.jiangxi.gov.cn/"),
    ("南昌市人社局", "http://rsj.nc.gov.cn/"),
    ("江西省国资委", "http://gzw.jiangxi.gov.cn/"),
    ("南昌市国资委", "http://gzw.nc.gov.cn/"),
    ("江西省住建厅", "http://zjt.jiangxi.gov.cn/"),
    ("南昌市住建局", "http://zjj.nc.gov.cn/"),
    ("南昌轨道交通集团", "http://www.ncmtr.com/"),
    ("江西省水利厅", "http://slt.jiangxi.gov.cn/"),
    ("赣州市人社局", "http://rsj.ganzhou.gov.cn/"),
    ("九江市人社局", "http://rsj.jiujiang.gov.cn/"),
]

NATIONAL_SITES = [
    ("国资委", "http://www.sasac.gov.cn/n2588025/n2588139/index.html"),
    ("住建部", "https://www.mohurd.gov.cn/"),
    ("水利部", "http://www.mwr.gov.cn/"),
    ("国聘网", "https://www.iguopin.com/"),
]

# ============================================================
#  Main
# ============================================================
def run(batch_name="manual"):
    log("=== 国企招聘采集 v4: %s ===" % batch_name)
    log("User: %dyo / %s / %s / %s / %s" % (
        USER_PROFILE["age"], USER_PROFILE["edu"], USER_PROFILE["major"],
        "/".join(USER_PROFILE["certs"]), USER_PROFILE["city"]))
    history = load_history()
    os.makedirs(TMP_DIR, exist_ok=True)
    all_found = []
    processed = []

    # Round 1: WeChat public accounts (relaxed matching - these are recruitment-focused)
    log("Round 1: WeChat public accounts")
    for acc in WECHAT_ACCOUNTS:
        log("  Searching: %s" % acc)
        articles = search_wechat_articles(acc, 1)
        for a in articles:
            if is_seen(a["title"], a["source"], history):
                continue
            # Known recruitment accounts: accept if it says recruitment
            if is_recruitment(a["title"]):
                all_found.append(a)
                log("  Hit: %s" % a["title"][:50])

    # Round 2: Jiangxi/Nanchang gov sites
    for name, url in LOCAL_GOV_SITES:
        if not url:
            continue
        log("  Scraping: %s" % name)
        jobs = scrape_government_site(name, url)
        for j in jobs:
            if is_seen(j["title"], j["source"], history):
                continue
            if is_eligible(j["title"]):
                all_found.append(j)
                log("  Hit: %s" % j["title"][:50])

    # Process if found
    if all_found:
        log("Found %d locally, processing..." % len(all_found))
        for job in all_found:
            result = process_job(job, history, TMP_DIR)
            processed.append(result)
        out_dir = make_output_dir()
        os.makedirs(out_dir, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        zip_name = "guopin_%s_%s.zip" % (batch_name, ts)
        zip_path = os.path.join(out_dir, zip_name)
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for result in processed:
                if os.path.exists(result["docx"]):
                    zf.write(result["docx"], "%s_article.docx" % result["title"][:20])
                for att_path in result["attachments"]:
                    if os.path.exists(att_path):
                        zf.write(att_path, "%s_%s" % (result["title"][:20],
                                                       os.path.basename(att_path)))
        log("Done: %s (%d items)" % (zip_path, len(processed)))
        shutil.rmtree(TMP_DIR, ignore_errors=True)
        save_history(history)
        return zip_path, len(processed), True

    # Round 3: National
    log("Round 3: National scope (local empty)")
    for name, url in NATIONAL_SITES:
        log("  Scraping: %s" % name)
        jobs = scrape_government_site(name, url)
        for j in jobs:
            if is_seen(j["title"], j["source"], history):
                continue
            if is_eligible(j["title"]):
                all_found.append(j)
                log("  Hit: %s" % j["title"][:50])

    if not all_found:
        log("No matching jobs today. Nothing created.")
        save_history(history)
        return None, 0, False

    log("Found %d nationally, processing..." % len(all_found))
    for job in all_found:
        result = process_job(job, history, TMP_DIR)
        processed.append(result)
    out_dir = make_output_dir()
    os.makedirs(out_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    zip_name = "guopin_national_%s_%s.zip" % (batch_name, ts)
    zip_path = os.path.join(out_dir, zip_name)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for result in processed:
            if os.path.exists(result["docx"]):
                zf.write(result["docx"], "%s_article.docx" % result["title"][:20])
            for att_path in result["attachments"]:
                if os.path.exists(att_path):
                    zf.write(att_path, "%s_%s" % (result["title"][:20],
                                                   os.path.basename(att_path)))
    log("Done: %s (%d items)" % (zip_path, len(processed)))
    shutil.rmtree(TMP_DIR, ignore_errors=True)
    save_history(history)
    return zip_path, len(processed), False

# ============================================================
#  Entry
# ============================================================
def fetch_zhaopin(keyword):
    return []

def fetch_51job(keyword):
    return []

if __name__ == "__main__":
    import sys
    batch = sys.argv[1] if len(sys.argv) > 1 else "manual"
    run(batch)
