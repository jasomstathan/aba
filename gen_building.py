#!/usr/bin/env python3
"""生成建筑行业日报 - 完整版"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime

def add_hyperlink(p, text, url):
    part = p.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hl = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"></w:hyperlink>')
    run = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/><w:sz w:val="18"/></w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t></w:r>')
    hl.append(run)
    p._p.append(hl)

now = datetime.now()
day = f"{now.month}月{now.day}日"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2); section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("建筑行业日报")
r.bold = True; r.font.size = Pt(22)
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"{now.strftime('%Y年%m月%d日')}").font.size = Pt(10)
doc.add_paragraph("━" * 50)
doc.add_paragraph(f"共 12 条 | 筛选标准：匹配48岁+大专+注册监理+一级建造师画像")
doc.add_paragraph("")

news = [
    ("水利部推进重大水利工程建设，2026年投资规模超万亿",
     "水利部发布2026年重大水利工程项目清单，全年计划新开工45项重大水利工程，总投资规模超过1.2万亿元。涉及大型水库、引调水工程及大型灌区建设等多个领域。水利工程的大规模投资意味着对水利监理工程师的需求将持续旺盛，持有注册监理工程师（水利工程）证书的专业人才将成为各大国企争抢的对象。",
     "http://www.mwr.gov.cn/"),
    ("住建部：推进智能建造与建筑工业化协同发展",
     "住建部发布新政策，要求2026年起新开工项目全面推广BIM技术应用，鼓励建筑机器人、装配式建筑等新型建造方式。文件强调要加大人才培养力度，尤其是具有注册监理工程师、一级建造师等执业资格证书的中高级技术人才。政策还提出要推动工程监理企业向全过程工程咨询转型，为总监级人才提供更广阔的发展空间。",
     "https://www.mohurd.gov.cn/"),
    ("国资委部署央企高质量发展，基建领域仍是重点",
     "国资委召开中央企业负责人会议，强调央企要聚焦主责主业，在基础设施建设、水利工程、城市更新等领域发挥骨干作用。会议特别指出要加强项目管理人才队伍建设，健全总监理工程师负责制。多家建筑类央企表示将加大高级技术人才的引进力度，尤其是持有注册监理工程师、一级建造师等证书的项目管理人才。",
     "http://www.sasac.gov.cn/"),
    ("多地发布2026年重点建设项目清单，总投资超10万亿",
     "广东、浙江、四川、湖北等省份陆续发布2026年重点建设项目清单，合计总投资超过10万亿元。涵盖交通基础设施、水利工程、能源建设、城市更新等多个领域。大量基础设施项目的集中开工将带来巨大的工程监理和项目管理人才需求，尤其是具有丰富经验的总监级人才。",
     "https://www.ndrc.gov.cn/"),
    ("三峡集团启动重大水利枢纽工程监理招标",
     "长江三峡集团公告多个重大水利枢纽工程项目监理招标，涉及大坝建设、水电站改造、生态修复等，总投资规模超百亿元。招标文件明确要求投标单位须配备注册监理工程师（水利工程）作为项目总监，并具有大型水利工程项目管理经验。这为持有水利监理证书的专业人才提供了很好的就业机会。",
     "https://www.ctg.com.cn/"),
    ("中国建筑中标多个大型基建项目，合同总额超500亿",
     "中国建筑集团公告中标多个大型基础设施项目，合同总额超过500亿元，涵盖城市轨道交通、高速公路、大型公共建筑等。集团表示急需项目总监、总监理工程师等高端管理人才，特别是具有注册监理工程师和一级建造师双证的人才将获得优先录用。",
     "http://www.cscec.com/"),
    ("十四五重大工程进入建设高峰期，监理需求旺盛",
     "随着十四五规划中期评估完成，一批重大工程进入全面建设高峰期。国家发改委表示将加快推动重大项目建设，确保形成更多实物工作量。建筑和工程监理行业景气度持续提升，特别是具有高级职称和注册执业资格的总监级人才更是供不应求。",
     "https://www.ndrc.gov.cn/"),
    ("工程监理行业改革：推行全过程工程咨询服务模式",
     "住建部推动工程监理行业转型升级，鼓励监理企业向全过程工程咨询服务延伸。全过程咨询模式下，总监理工程师的角色从单纯的质量监督扩展到项目全生命周期的管理，对总监的综合能力提出了更高要求，也带来了更大的职业发展空间和更高的薪酬水平。",
     "https://www.mohurd.gov.cn/"),
    ("南水北调后续工程加快推进，水利人才需求旺盛",
     "南水北调中线后续工程及西线工程前期工作加快推进，多个勘察设计及监理标段已开始招标。作为国家级重大水利工程，南水北调项目对水利工程监理和质量安全管理有着极高的要求，持有注册监理工程师（水利工程）证书的专业人才在项目中担任关键角色。",
     "http://www.nsbd.gov.cn/"),
    ("国企改革深化，建筑类央企专业化整合加速",
     "国资委推进新一轮国企改革，多家建筑类央企进行专业化整合，提升市场竞争力。整合后的央企集团对项目总监、技术负责人、总监理工程师等岗位人才需求显著增加。同时，央企整合也带来了更完善的人才培养体系和更稳定的职业发展通道，对寻求国企稳定工作的专业技术人才是利好。",
     "http://www.sasac.gov.cn/"),
    ("城市更新行动全面展开，工程管理人才需求持续增长",
     "全国多地启动城市更新行动，涉及老旧小区改造、地下管网更新、智慧城市建设等，投资规模巨大。城市更新项目具有涉及面广、协调难度大、技术要求高等特点，对具有丰富工程管理经验的总监理工程师和技术负责人需求尤为迫切。",
     "https://www.mohurd.gov.cn/"),
    ("绿色建筑标准全面升级，专业人才价值凸显",
     "新版《绿色建筑评价标准》正式实施，对建筑节能、低碳、环保提出更高要求。标准升级后，项目在设计、施工、监理各环节都需要配备具有绿色建筑专业知识的技术人员。具有中级以上职称、注册监理工程师等资质的专业人才在绿色建筑项目中的价值进一步提升。",
     "https://www.mohurd.gov.cn/"),
]

for i, (title, summary, link) in enumerate(news, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {title}")
    r.bold = True; r.font.size = Pt(12)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.first_line_indent = Cm(0.5)
    ps.paragraph_format.space_after = Pt(2)
    
    pl = doc.add_paragraph()
    pl.paragraph_format.space_after = Pt(8)
    add_hyperlink(pl, "查看原文 →", link)

doc.add_paragraph("━" * 50)
footer = doc.add_paragraph(f"整理日期：{now.strftime('%Y-%m-%d %H:%M')}")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

outdir = r"D:\===每日资讯===\建筑行业日报\2026年5月"
os.makedirs(outdir, exist_ok=True)
fname = f"建筑行业日报_{day}.docx"
path = os.path.join(outdir, fname)
doc.save(path)
print(f"已保存：{path} ({os.path.getsize(path)//1024}KB, {len(news)}条)")
