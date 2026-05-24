#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========= Page setup =========
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ========= Styles =========
style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(14)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
style.paragraph_format.line_spacing = 1.5

# ========= Helper =========
def add_run(paragraph, text, bold=False, size=None, font=None, color=None, underline=False):
    run = paragraph.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if font: run.font.name = font; run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color: run.font.color.rgb = RGBColor(*color)
    if underline: run.font.underline = True
    return run

def new_para(align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

# ====== Title ======
t = new_para(WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_run(t, '工程监理通知单', bold=True, size=22, font='黑体')

# ====== Subtitle ======
sub = new_para(WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_run(sub, '（安全生产类）', size=14, font='仿宋')

# ====== Meta info ======
meta = new_para(space_before=6, space_after=6)
add_run(meta, '编号：', bold=True, size=14, font='黑体')
add_run(meta, 'JL-TZ-2026-005', size=14, font='仿宋')

meta2 = new_para(space_after=6)
add_run(meta2, '致：', bold=True, size=14, font='黑体')
add_run(meta2, '（施工单位名称）', size=14, font='仿宋', underline=True)

# ====== Subject ======
subj = new_para(space_before=4, space_after=4)
add_run(subj, '事由：', bold=True, size=14, font='黑体')
add_run(subj, '施工现场重大安全隐患整改通知', size=14, font='仿宋')

# ====== Content ======
content = new_para(first_line_indent=0.74)
add_run(content, '内容：', bold=True, size=14, font='黑体')

p1 = new_para(first_line_indent=0.74)
add_run(p1, '我监理部于 2026 年 5 月 23 日对施工现场进行安全专项检查，发现存在以下严重安全隐患：', size=14, font='仿宋')

# ---- Section 1 ----
s1 = new_para(space_before=4)
add_run(s1, '一、高处作业及临边防护方面', bold=True, size=14, font='黑体')

items1 = [
    '高处作业人员未系安全带，存在高坠重大风险；',
    '移动作业架搭设不规范，缺少防护栏杆，作业人员无安全立足点；',
    '现场违规使用木质人字梯进行高处作业，刚度及稳定性不满足安全要求。',
]
for item in items1:
    ip = new_para(first_line_indent=0.74)
    add_run(ip, f'{item}', size=14, font='仿宋')

# ---- Section 2 ----
s2 = new_para(space_before=4)
add_run(s2, '二、人员安全防护方面', bold=True, size=14, font='黑体')

items2 = [
    '施工现场部分作业人员未按规定佩戴安全帽；',
    '现场电缆线随意拖地敷设，未采取过路保护措施，存在触电及绊倒隐患。',
]
for item in items2:
    ip = new_para(first_line_indent=0.74)
    add_run(ip, f'{item}', size=14, font='仿宋')

# ---- Section 3 ----
s3 = new_para(space_before=4)
add_run(s3, '三、动火作业管理方面', bold=True, size=14, font='黑体')

items3 = [
    '食堂切割作业区动火作业许可证已过期，仍继续进行切割作业；',
    '报告厅焊接作业未设置接火斗，焊渣散落，极易引发火灾；',
    '切割作业区未设置挡火板，火花飞溅无有效隔离。',
]
for item in items3:
    ip = new_para(first_line_indent=0.74)
    add_run(ip, f'{item}', size=14, font='仿宋')

# ---- Regulation reference ----
ref = new_para(space_before=4, first_line_indent=0.74)
add_run(ref, '上述问题已严重违反《建设工程安全生产管理条例》《建筑施工高处作业安全技术规范》（JGJ 80-2016）及《施工现场消防安全技术规范》（GB 50720-2011）相关规定。', size=14, font='仿宋')

# ====== Order ======
order_title = new_para(space_before=6)
add_run(order_title, '现责令你部：', bold=True, size=14, font='黑体')

orders = [
    '自收到本通知之日起立即停止所有高处及动火作业；',
    '对上述隐患逐条整改，于 24 小时内整改完毕；',
    '对所有高处作业人员进行重新安全交底，不合格者严禁上岗；',
    '动火作业许可证未办妥或过期未续的，严禁进行任何动火作业；',
    '整改完成后书面回复监理部，经复查合格后方可复工。',
]
for i, order in enumerate(orders, 1):
    op = new_para(first_line_indent=0.74)
    add_run(op, f'{i}. {order}', size=14, font='仿宋')

# ====== Warning ======
warn = new_para(space_before=6, first_line_indent=0.74)
add_run(warn, '若逾期未整改或整改不到位，我部将依据合同条款下发工程暂停令，并上报建设行政主管部门。由此造成的工期延误及一切损失由贵部自行承担。', bold=True, size=14, font='仿宋', color=(204, 0, 0))

# ====== Signatures ======
sig1 = new_para(space_before=18)
add_run(sig1, '监理工程师（签字）：', size=14, font='仿宋')

sig2 = new_para(space_before=6)
add_run(sig2, '总监理工程师（签字）：', size=14, font='仿宋')

sig3 = new_para(space_before=6)
add_run(sig3, '（项目监理部盖章）', size=14, font='仿宋')

date_p = new_para(space_before=12)
add_run(date_p, '日期：2026 年 5 月 23 日', size=14, font='仿宋')

# ====== CC ======
cc = new_para(space_before=14)
add_run(cc, '抄送：建设单位', size=14, font='仿宋')

# ====== Save ======
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
output = os.path.join(desktop, '监理处罚通知单_安全隐患整改.docx')
doc.save(output)
print(f'文档已保存到：{output}')
