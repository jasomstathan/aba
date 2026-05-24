#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""多源资讯采集器 v4 - 仅保留有内容的网站"""
import os, re, json, hashlib, ssl, urllib.request, sys
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
HDR = {"User-Agent": "Mozilla/5.0"}
BASE_DIR = r"D:\===每日资讯==="
HISTORY_FILE = os.path.join(BASE_DIR, "_history.json")

def load_history():
    try: return json.load(open(HISTORY_FILE, "r", encoding="utf-8"))
    except: return {"articles": []}

def save_history(h):
    cutoff = (datetime.now()-timedelta(days=60)).strftime("%Y-%m-%d")
    h["articles"] = [a for a in h["articles"] if a.get("date","") >= cutoff]
    json.dump(h, open(HISTORY_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=2)

def key_of(t): return hashlib.md5(re.sub(r"[^\u4e00-\u9fa5a-zA-Z0-9]","",t.lower())[:25].encode()).hexdigest()

def is_dup(t, h, seen):
    k = key_of(t)
    if k in seen: return True
    for a in h["articles"]:
        if a.get("key")==k and (datetime.now()-datetime.strptime(a["date"],"%Y-%m-%d")).days<7: return True
    return False

def mark(t,s,cat,h,seen):
    k=key_of(t); seen.add(k)
    h["articles"].append({"key":k,"title":t[:40],"category":cat,"source":s,"date":datetime.now().strftime("%Y-%m-%d")})

def fetch(url):
    try:
        r = urllib.request.urlopen(urllib.request.Request(url,headers=HDR), context=ctx, timeout=10)
        return r.read().decode("utf-8","replace")
    except: return ""

def extract_items(html):
    items = []
    blocks = re.split(r'</?(?:div|li|p|section|article|tr|td)[^>]*>', html)
    for block in blocks:
        for a_html in re.findall(r'<a[^>]*>(.+?)</a>', block, re.DOTALL):
            text = re.sub(r'<[^>]+>', '', a_html).strip()
            if len(text) < 8: continue
            hm = re.search(r'href="([^"]+)"', a_html) or re.search(r"href='([^']+)'", a_html)
            if not hm: hm = re.search(r'href="([^"]+)"', block)
            link = hm.group(1) if hm else ""
            if not link or not link.startswith("http"): continue
            summ = block.replace(a_html, "")
            summ = re.sub(r'<[^>]+>', '', summ).strip()
            summ = re.sub(r'\s+', ' ', summ)[:200]
            if not summ or len(summ) < 20 or summ == text: summ = ""
            items.append({"title": text, "link": link, "summary": summ})
    seen_t = set()
    return [it for it in items if not (it["title"][:15] in seen_t or seen_t.add(it["title"][:15]))]

# ===== 只有能出摘要/内容的网站 =====
AI_S = [
    ("OSCHINA", "https://www.oschina.net/news/industry"),   # 有摘要
    ("雷锋网", "https://www.leiphone.com/category/ai"),     # 有摘要
    ("新浪科技", "https://tech.sina.com.cn/"),              # 有摘要
    ("网易科技", "https://tech.163.com/"),                  # 部分有
    ("百度新闻", "https://news.baidu.com/"),                # 有
    ("工信部", "https://www.miit.gov.cn/"),                 # 有
    ("国家能源局", "http://www.nea.gov.cn/"),               # 有
    ("中国科技网", "http://www.stdaily.com/"),              # 有
    ("新华网科技", "http://www.xinhuanet.com/tech/"),       # 有
    ("人民网科技", "http://scitech.people.com.cn/"),        # 有
    ("澎湃科技", "https://www.thepaper.cn/"),               # 有
]

BLD_S = [
    ("住建部", "https://www.mohurd.gov.cn/"),               # 有
    ("国资委", "http://www.sasac.gov.cn/"),                 # 有
    ("发改委", "https://www.ndrc.gov.cn/"),                 # 有
    ("水利部", "http://www.mwr.gov.cn/"),                   # 有
    ("交通运输部", "https://www.mot.gov.cn/"),              # 有
    ("自然资源部", "https://www.mnr.gov.cn/"),             # 有
    ("国家能源局", "http://www.nea.gov.cn/"),               # 有
    ("OSCHINA", "https://www.oschina.net/news/industry"),   # 有
    ("新浪科技", "https://tech.sina.com.cn/"),              # 有
    ("百度新闻", "https://news.baidu.com/"),                # 有
    ("网易科技", "https://tech.163.com/"),                  # 部分有
    ("中国政府采购网", "http://www.ccgp.gov.cn/"),          # 有
]

# ===== 搜狗微信搜索（专门搜公众号文章）=====
def search_wechat(keyword, page=1):
    """搜狗微信搜索公众号文章"""
    import urllib.parse
    kw = urllib.parse.quote(keyword)
    url = f"https://weixin.sogou.com/weixin?type=2&query={kw}&page={page}&ie=utf8"
    html = fetch(url)
    if not html: return []
    
    items = []
    # 提取文章标题和链接
    for m in re.finditer(r'<h3[^>]*>\s*<a[^>]*href="([^"]+)"[^>]*>\s*<em[^>]*>([^<]+)</em>\s*</a>', html):
        link = m.group(1)
        title = re.sub(r'<[^>]+>', '', m.group(2)).strip()
        if link.startswith("/"): link = "https://weixin.sogou.com" + link
        items.append({"title": title, "link": link, "source": "微信公众号"})
    
    return items

def collect(cat, sources, h, seen):
    arts = []
    
    # 先搜微信公众号（AI早报专用）
    if cat == "AI科技早报":
        wx_kws = ["AI技术分享","AI工具评测","人工智能干货","ChatGPT教程","AI实用技巧"]
        for kw in wx_kws:
            for it in search_wechat(kw):
                if is_dup(it["title"], h, seen): continue
                if not any(k in it["title"] for k in ["AI","人工智能","GPT","模型","智能","机器","数据","编程","工具","教程","技巧","干货","体验","评测","分享","实战"]): continue
                mark(it["title"], it["source"], cat, h, seen)
                arts.append(it)
                if len(arts) >= 15: break
            if len(arts) >= 15: break
    
    # 再从网站采集
    for name, url in sources:
        if len([a for a in arts if a.get("summary")]) >= 35: break
        html = fetch(url)
        if not html or len(html) < 500: continue
        for it in extract_items(html):
            if is_dup(it["title"], h, seen): continue
            kws = (["AI","人工智能","大模型","GPT","Chat","机器人","编程","数据","算法","芯片","开源","科技","智能","互联网","数字","系统","软件","工具","教程","技巧","分享","实战","评测"]
                   if cat == "AI科技早报" else
                   ["建筑","工程","施工","监理","建设","水利","土木","桥梁","招标","项目","规划","安全","质量","基础设施","能源","交通","城市","改造","地产","材料","设计"])
            if not any(k.lower() in it["title"].lower() for k in kws): continue
            mark(it["title"], name, cat, h, seen)
            it["source"] = name
            arts.append(it)
    
    # 去重排序：有摘要的优先
    arts.sort(key=lambda x: (1 if x.get("summary") else 0, x["title"]), reverse=True)
    return arts[:50]

def add_link(p, text, url):
    part=p.part; rid=part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    hl=parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{rid}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"></w:hyperlink>')
    run=parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/><w:sz w:val="20"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>')
    hl.append(run); p._p.append(hl)

def gen_docx(cat, arts, day):
    doc=Document(); s=doc.sections[0]; s.top_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    st=doc.styles["Normal"]; st.font.name="宋体"; st.font.size=Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体")
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run(cat); r.bold=True; r.font.size=Pt(22); r.font.name="黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体"); r.font.color.rgb=RGBColor(0x1a,0x3c,0x6e)
    doc.add_paragraph("━"*50)
    doc.add_paragraph(f"共 {len(arts)} 条 | 来源：{'/'.join(set(a['source'] for a in arts))}")
    doc.add_paragraph("")
    for i,a in enumerate(arts,1):
        p=doc.add_paragraph(); r=p.add_run(f"{i}. {a['title']}"); r.bold=True; r.font.size=Pt(11)
        r.font.name="黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体")
        if a.get("summary"):
            ps=doc.add_paragraph(a["summary"]); ps.paragraph_format.first_line_indent=Cm(0.5); ps.paragraph_format.space_after=Pt(2)
        pl=doc.add_paragraph(); pl.paragraph_format.space_after=Pt(6)
        pl.add_run(f"[{a['source']}]  ").font.size=Pt(8)
        add_link(pl,"查看原文→",a["link"])
    doc.add_paragraph("━"*50)
    f=doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"); f.alignment=WD_ALIGN_PARAGRAPH.CENTER
    return doc

def run(cat):
    now=datetime.now(); day=f"{now.month}月{now.day}日"
    h=load_history(); seen=set()
    srcs=AI_S if cat=="AI科技早报" else BLD_S
    outdir=os.path.join(BASE_DIR,cat,f"{now.year}年{now.month}月"); os.makedirs(outdir,exist_ok=True)
    print(f"采集 [{cat}] - {len(srcs)}个有内容来源")
    arts=collect(cat,srcs,h,seen)
    # 没有摘要的条目只保留标题
    print(f"收录：{len(arts)}条（{len(set(a['source'] for a in arts))}个网站），其中有摘要：{len([a for a in arts if a.get('summary')])}条")
    if not arts:
        print("无新内容")
        save_history(h)
        return
    doc=gen_docx(cat,arts,day)
    fname=f"{cat}_{day}.docx"; path=os.path.join(outdir,fname)
    try: doc.save(path)
    except PermissionError: doc.save(os.path.join(outdir,f"_{fname}"))
    print(f"已保存：{path}")
    save_history(h)

if __name__=="__main__":
    run(sys.argv[1] if len(sys.argv)>1 else "AI科技早报")
