#!/usr/bin/env python3
"""生成晚间复盘 - 完整版"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

now = datetime.now()
day = f"{now.month}月{now.day}日"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2); section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("晚间复盘")
r.bold = True; r.font.size = Pt(22)
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"{now.strftime('%Y年%m月%d日')} 星期日").font.size = Pt(10)
doc.add_paragraph("━" * 40)

doc.add_heading("今日回顾", level=2)
items = [
    "完成了OpenClaw AI助手的全面配置和调试工作",
    "安装了Claude Code、CC-Switch、OpenAI Codex等AI编程工具",
    "搭建了国企招聘信息自动采集系统，每日定时采集并生成Word报告",
    "创建了每日资讯自动归档系统，包括AI科技早报和建筑行业日报",
    "整理并测试了所有定时任务，确认各模块正常运行",
    "浏览了智联招聘上的建筑行业岗位信息，重点关注总监级职位",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("招聘匹配分析", level=2)
doc.add_paragraph("当前筛选条件：48岁 + 大专 + 工民建专业 + 中级职称")
doc.add_paragraph("证书资质：注册监理工程师（建筑专业+水利专业） + 一级建造师（建筑专业）")
doc.add_paragraph("")
doc.add_paragraph("匹配方向分析：")
matches = [
    "总监理工程师/项目总监 — 年龄放宽到50-55岁，大专可接受，匹配度高",
    "水利工程总监/总监代表 — 水利监理证书稀缺，国企需求大，匹配度高",
    "技术负责人/总工程师 — 需要高级职称，可作为长期目标",
    "工程部经理/项目经理 — 一建证书硬通货，但需关注学历要求",
]
for m in matches:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading("明日计划", level=2)
plans = [
    "关注各招聘网站新增的监理总监级岗位",
    "查看今日AI科技早报和建筑行业日报中的行业动态",
    "如有合适的岗位，准备投递材料",
]
for p in plans:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph("")
doc.add_paragraph("━" * 40)
note = doc.add_paragraph("每天进步一点点，坚持带来大改变。")
note.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

outdir = r"D:\===每日资讯===\晚间复盘\2026年5月"
os.makedirs(outdir, exist_ok=True)
fname = f"晚间复盘_{day}.docx"
path = os.path.join(outdir, fname)
doc.save(path)
print(f"已保存：{path} ({os.path.getsize(path)//1024}KB)")
