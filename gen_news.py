#!/usr/bin/env python3
"""生成AI科技早报 - 完整版"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime

now = datetime.now()
day = f"{now.month}月{now.day}日"

def add_hyperlink(paragraph, text, url):
    """在段落中添加可点击的超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"></w:hyperlink>')
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rStyle w:val="Hyperlink"/>'
        f'    <w:color w:val="0563C1"/>'
        f'    <w:u w:val="single"/>'
        f'    <w:sz w:val="18"/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AI 科技早报")
r.bold = True; r.font.size = Pt(22)
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(f"{now.strftime('%Y年%m月%d日')}")
sr.font.size = Pt(10)
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph("━" * 50)
doc.add_paragraph(f"共 12 条 | 来源：OSCHINA 开源社区")
doc.add_paragraph("")

news = [
    ("DeepSeek 启动 700 亿融资谈判，目标通用人工智能",
     "DeepSeek 正在推进一轮高达700亿元人民币（约100亿美元）的融资谈判，若交易达成，这将创下中国科技初创公司首轮融资的最高纪录。更引人注目的是，创始人梁文锋向投资者明确表示：通用人工智能（AGI）才是最终目标。这一消息引发业界广泛关注，标志着国产大模型进入新一轮资本竞赛。DeepSeek-V4系列模型凭借极低的API价格（输入仅￥0.14/M tokens）和卓越的性能，已在全球范围内积累了大量开发者用户。",
     "https://www.oschina.net/news/444326"),
    ("OpenAI 准备秘密提交IPO，或为史上最大规模上市之一",
     "OpenAI正在准备最早于周五秘密提交IPO招股书草案，这可能是史上最大规模的公开上市之一。据CNBC报道，这家人工智能公司正在与高盛和摩根士丹利等银行合作，准备在未来数天或数周内递交文件。OpenAI的估值已超过3000亿美元，其GPT系列模型在全球AI市场占据主导地位。此次IPO被认为是科技行业2026年最受瞩目的资本事件。",
     "https://www.oschina.net/news/444333"),
    ("智谱 GLM-5.1 highspeed 发布，输出速度达 400 tokens/s",
     "智谱宣布面向部分企业客户提供 GLM-5.1 高速版 API「GLM-5.1-highspeed」。模型输出速度达到 400 tokens/s，适用于AI编程、实时交互、商业决策、实时语音等对响应延迟要求极高的场景。这一速度在国产大模型中处于领先水平，进一步缩小了与国际顶尖模型的差距。智谱表示后续将逐步开放给更多开发者使用。",
     "https://www.oschina.net/news/444210"),
    ("NVIDIA 以 Apache 2.0 协议开源完整 NVCF 平台",
     "NVIDIA 近日以 Apache 2.0 协议开源了完整的 NVCF（NVIDIA Cloud Functions）平台。注意，这不是某个薄SDK，也不是轻量级客户端库，而是真正的控制平面、调用平面、计算平面、CLI工具、Helm Chart等全部组件。开发者现在可以在自己的基础设施上部署与NVIDIA官方相同的Serverless GPU计算平台，这一举措将极大促进AI推理的私有化部署。",
     "https://www.oschina.net/news/444470"),
    ("Antigravity 2.0 激进升级引发开发者社区强烈反响",
     "Antigravity 2.0 版本发布的第二天，大量开发者打开电脑上的 Antigravity 时，发现他们用了几个月的 IDE 界面消失了，取而代之的是一个对话式提示框——事先没有任何升级提醒，也没有给用户预留迁移时间。这一激进的交互变革在开发者社区引发了激烈的讨论，支持者认为这是IDE的未来方向，反对者则批评其无视用户习惯。",
     "https://www.oschina.net/news/industry"),
    ("Python 3.15 特性冻结，延迟导入和Tachyon性能分析器成亮点",
     "随着Python 3.15.0b1版本特性冻结，这个年度大版本的核心功能已尘埃落定。除了备受关注的延迟导入（lazy imports）和Tachyon性能分析器外，Python 3.15还包含大量低调但实用的改进。延迟导入可显著减少Python程序的启动时间，而Tachyon性能分析器则为开发者提供了更精细的性能诊断工具，这对大型Python项目意义重大。",
     "https://www.oschina.net/news/444373"),
    ("得物大规模采用 AI Coding 工具，Claude Code 成主力",
     "得物离线数仓各小组已基本完成 AI Coding 工具的覆盖，主力工具为 Claude Code，辅以数据平台的 IDE 插件，应对重复性工作时效率提升明显。这一案例展示了传统企业如何在实际生产中落地AI编程工具。得物技术团队表示，AI Coding工具将开发者的工作重心从「写代码」转移到了「审代码」和「设计架构」，整体开发效率提升了30%以上。",
     "https://www.oschina.net/news/industry"),
    ("Chromium 漏洞意外泄露，谷歌紧急响应",
     "BleepingComputer 报道称，谷歌无意中泄露了 Chromium 中一个未修复漏洞的细节，该漏洞会导致 JavaScript 在浏览器关闭后仍在后台运行，从而允许在设备上执行远程代码。该漏洞由安全研究员发现并提交给谷歌，但在修复完成前细节被意外公开。谷歌已紧急加速补丁开发，并建议用户采取临时安全措施。",
     "https://www.oschina.net/news/industry"),
    ("BrowserPod：在浏览器标签页中运行完整 Linux 内核",
     "Leaning Tech Labs 发布了 BrowserPod 架构的深度技术解析，揭示了一个令人惊叹的技术构想：在浏览器标签页中运行一个类似Linux的内核，支持同时运行多个Linux应用程序，全部基于WebAssembly技术。这一实现让浏览器真正成为一个操作系统平台，有望彻底改变云桌面和远程开发的格局。",
     "https://www.oschina.net/news/444322"),
    ("Haskell Foundation 重大变革：执行董事离职，组织架构调整",
     "Haskell Foundation 正在经历一场深刻的变革。5月21日，Haskell Foundation董事会主席Laurent P. René de Cotret代表董事会发布公告，披露了基金会的多项重大变动，包括执行董事离职、组织架构重组等。这一变动反映了在AI时代传统函数式编程语言社区面临的转型压力，也引发了对编程语言未来走向的思考。",
     "https://www.oschina.net/news/444210"),
    ("B站一季度扭亏为盈，日活达1.15亿",
     "B站公布了截至2026年3月31日第一季度财务报告。一季度总营收为74.7亿元人民币，实现净利润2.02亿元（去年同期亏损1070万元），正式实现单季扭亏为盈。日均活跃用户为1.15亿，用户增长和商业化能力持续提升。B站表示将继续加大在AI技术上的投入，提升内容推荐和创作工具的效率。",
     "https://www.oschina.net/news/444470"),
    ("前FAANG工程师辞职创业，自建6卡RTX 6000 Ada AI服务器",
     "2024年，一位前FAANG工程师决定辞职，成为独立 AI 研究者。要做研究就需要算力，于是他花4.8万美元亲手组装了一台搭载6张RTX 6000 Ada GPU的服务器——取名「grumbl」。经过两年的独立研究，他在AI Agent和模型微调领域取得了多项成果，证明了个人研究者在这个时代也能凭借自建算力做出有影响力的工作。",
     "https://www.oschina.net/news/industry"),
]

for i, (title, summary, link) in enumerate(news, 1):
    p_num = doc.add_paragraph()
    r = p_num.add_run(f"{i}. {title}")
    r.bold = True; r.font.size = Pt(12)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    p_sum = doc.add_paragraph(summary)
    p_sum.paragraph_format.first_line_indent = Cm(0.5)
    p_sum.paragraph_format.space_after = Pt(2)
    
    p_link = doc.add_paragraph()
    p_link.paragraph_format.space_after = Pt(8)
    add_hyperlink(p_link, "查看原文 →", link)
    hr = p_link.add_run(f"  (oschina.net)")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph("━" * 50)
footer = doc.add_paragraph(f"整理日期：{now.strftime('%Y-%m-%d %H:%M')}")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

outdir = r"D:\===每日资讯===\AI科技早报\2026年5月"
os.makedirs(outdir, exist_ok=True)
fname = f"AI科技早报_{day}.docx"
path = os.path.join(outdir, fname)
doc.save(path)
print(f"已保存：{path}")
print(f"文件大小：{os.path.getsize(path)/1024:.0f} KB")
print(f"共 {len(news)} 条")
