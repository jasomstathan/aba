#!/usr/bin/env python3
"""生成晚间复盘Word文档 - 测试运行"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

now = datetime.now()
day = f"{now.month}月{now.day}日"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("晚间复盘")
r.bold = True; r.font.size = Pt(22)
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"{now.strftime('%Y年%m月%d日')}").font.size = Pt(10)
doc.add_paragraph("━" * 40)

doc.add_heading("今日回顾", level=2)
doc.add_paragraph("• 查看并梳理了国企招聘信息，筛选总监级岗位")
doc.add_paragraph("• 了解建筑行业最新政策动态")
doc.add_paragraph("• 整理了AI科技资讯")

doc.add_heading("招聘进展", level=2)
doc.add_paragraph("• 智联招聘上搜到多个总监级岗位")
doc.add_paragraph("• 重点关注：水利工程总监、总监理工程师等职位")
doc.add_paragraph("• 匹配条件：48岁+大专+注册监理(建筑+水利)+一级建造师")

doc.add_heading("明日计划", level=2)
doc.add_paragraph("• 继续关注国企招聘信息发布")
doc.add_paragraph("• 查看今日收集的岗位详情，筛选适合的投递")
doc.add_paragraph("• 关注水利工程方向的新增岗位")

doc.add_paragraph("")
doc.add_paragraph("━" * 40)
note = doc.add_paragraph("睡前想一想：今天做了哪些事？明天要做的最重要的一件事是什么？")
note.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
note.runs[0].font.size = Pt(9)

outdir = r"D:\===每日资讯===\晚间复盘\2026年5月"
os.makedirs(outdir, exist_ok=True)
fname = f"晚间复盘_{day}.docx"
path = os.path.join(outdir, fname)
doc.save(path)
print(f"已保存：{path} ({os.path.getsize(path)} bytes)")
