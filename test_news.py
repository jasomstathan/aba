#!/usr/bin/env python3
"""生成AI科技早报Word文档 - 测试运行"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

now = datetime.now()
day = f"{now.month}月{now.day}日"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AI 科技早报")
r.bold = True; r.font.size = Pt(22)
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(f"{datetime.now().strftime('%Y年%m月%d日')}")
sr.font.size = Pt(10)
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
sr.font.name = '宋体'
sr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph("━" * 50)

# 新闻条目（从web_fetch数据提取）
news = [
    ("DeepSeek 启动 700 亿融资谈判", 
     "DeepSeek 正在推进约100亿美元的融资谈判，若达成将创下中国科技初创公司首轮融资最高纪录。创始人梁文锋明确表示AGI是目标。",
     "https://www.oschina.net/news/444326"),
    ("OpenAI 准备秘密提交IPO",
     "OpenAI最早将于本周秘密提交IPO招股书草案，与高盛和摩根士丹利合作，可能是史上最大规模上市之一。",
     "https://www.oschina.net/news/444333"),
    ("智谱 GLM-5.1 highspeed 发布",
     "智谱推出GLM-5.1高速版API，输出速度达400 tokens/s，面向AI编程、实时交互、商业决策等场景。",
     "https://www.oschina.net/news/444210"),
    ("NVIDIA 开源 NVCF 平台",
     "NVIDIA以Apache 2.0协议开源完整的NVCF（NVIDIA Cloud Functions）平台，包含控制平面、CLI工具等全部组件。",
     "https://www.oschina.net/news/444470"),
    ("Antigravity 2.0 激进升级引发争议",
     "大量开发者发现IDE界面消失，被对话式提示框取代。事先无升级提醒，引发社区广泛讨论。",
     "https://www.oschina.net/news/industry"),
    ("Python 3.15 特性冻结",
     "Python 3.15.0b1版本特性冻结，核心功能包括延迟导入（lazy imports）和Tachyon性能分析器。",
     "https://www.oschina.net/news/444373"),
    ("得物采用 Claude Code 提升编程效率",
     "得物离线数仓全面覆盖AI Coding工具，主力为Claude Code，效率提升明显。",
     "https://www.oschina.net/news/industry"),
    ("Chromium 漏洞泄露事件",
     "谷歌无意中泄露未修复漏洞细节，该漏洞导致JS在浏览器关闭后仍在后台运行，可执行远程代码。",
     "https://www.oschina.net/news/industry"),
    ("BrowserPod：在浏览器中运行Linux内核",
     "Leaning Tech Labs发布技术解析，在浏览器标签页中运行类Linux内核，全部基于WebAssembly技术。",
     "https://www.oschina.net/news/444322"),
    ("Haskell Foundation 重大变革",
     "Haskell Foundation执行董事离职、组织架构调整，AI时代下函数式编程语言的转型。",
     "https://www.oschina.net/news/444210"),
    ("B站一季度扭亏为盈",
     "B站一季度总营收74.7亿元，净利润2.02亿元，日活1.15亿，正式实现单季扭亏为盈。",
     "https://www.oschina.net/news/444470"),
    ("前FAANG工程师自建AI服务器",
     "一位前FAANG工程师辞职成为独立AI研究者，花4.8万美元组装6张RTX 6000 Ada GPU的服务器。",
     "https://www.oschina.net/news/industry"),
]

for i, (title, summary, link) in enumerate(news, 1):
    # 标题
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {title}")
    r.bold = True; r.font.size = Pt(12)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 摘要
    p2 = doc.add_paragraph(summary)
    p2.paragraph_format.first_line_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(1)
    
    # 链接
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run(f"原文链接：{link}")
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

doc.add_paragraph("━" * 50)
footer = doc.add_paragraph(f"来源：OSCHINA | 整理日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.runs[0]
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 保存
outdir = r"D:\===每日资讯===\AI科技早报\2026年5月"
os.makedirs(outdir, exist_ok=True)
fname = f"AI科技早报_{day}.docx"
path = os.path.join(outdir, fname)
doc.save(path)
print(f"已保存：{path} ({os.path.getsize(path)} bytes, {len(news)}条)")
